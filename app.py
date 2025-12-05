from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory, jsonify, g, send_file, flash, abort
from flask_wtf.csrf import CSRFProtect, generate_csrf
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from docx import Document
from openai import AzureOpenAI
from dotenv import load_dotenv
import os
import time
from datetime import datetime, timedelta
import requests
import markdown
from bs4 import BeautifulSoup
from flask_session import Session
import json
import pyodbc
from io import BytesIO
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import hashlib
import secrets
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
import base64
import io
from PIL import Image
import uuid
from urllib.parse import unquote
import asyncio
import httpx
import random
import aiohttp
import logging
from logging.handlers import RotatingFileHandler
from html import escape
from typing import Optional, Tuple, Dict, Any, List
from utils.validation import validate_email, validate_password, validate_text_input, sanitize_input

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        RotatingFileHandler('app.log', maxBytes=10485760, backupCount=10),  # 10MB per file, 10 backups
        logging.StreamHandler()  # Also log to console
    ]
)
logger = logging.getLogger(__name__)

# Set log level based on environment
if os.getenv('FLASK_DEBUG', 'false').lower() == 'true':
    logger.setLevel(logging.DEBUG)
else:
    logger.setLevel(logging.INFO)

load_dotenv()

app = Flask(__name__, static_folder='static')
app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config['UPLOAD_FOLDER'] = 'static/generated'

# CSRF Protection
csrf = CSRFProtect(app)
app.config['WTF_CSRF_ENABLED'] = True
app.config['WTF_CSRF_TIME_LIMIT'] = 3600  # 1 hour

# Make csrf_token available in all templates
@app.context_processor
def inject_csrf_token():
    """Make CSRF token function available in all templates"""
    def get_csrf_token():
        return generate_csrf()
    return dict(csrf_token=get_csrf_token)

# Rate Limiting
def get_user_id_or_ip() -> str:
    """
    Get user ID from session if logged in, otherwise use IP address.
    
    Used as key function for rate limiting to track limits per user or IP.
    
    Returns:
        User ID as string if logged in, otherwise IP address
    """
    try:
        if 'user' in session and session.get('user', {}).get('id'):
            return str(session['user']['id'])
    except RuntimeError:
        # Outside request context
        pass
    return get_remote_address()

limiter = Limiter(
    app=app,
    key_func=get_user_id_or_ip,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://",  # In-memory storage (use Redis in production)
    strategy="fixed-window"
)

# Database configuration
app.config['AZURE_SQL_SERVER'] = os.getenv('AZURE_SQL_SERVER')
app.config['AZURE_SQL_DATABASE'] = os.getenv('AZURE_SQL_DATABASE')
app.config['AZURE_SQL_USERNAME'] = os.getenv('AZURE_SQL_USERNAME')
app.config['AZURE_SQL_PASSWORD'] = os.getenv('AZURE_SQL_PASSWORD')

# Session configuration
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=1)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
FUNCTION_APP_URL = os.getenv("AZURE_FUNCTION_APP_URL")
FUNCTION_KEY = os.getenv("FUNCTION_KEY")
Session(app)

# Simulate OpenAI calls for testing
SIMULATE_OPENAI = os.getenv("SIMULATE_OPENAI", "false").lower() == "true"

async def simulate_openai_call():
    """Simulate OpenAI call with random delay"""
    delay = random.uniform(20, 40)  # Random delay between 20-40 seconds
    await asyncio.sleep(delay)
    return {"content": "Simulated response after delay"}


async def make_async_request(url, payload):
    """Make async HTTP request with error handling"""
    if SIMULATE_OPENAI:
        return await simulate_openai_call()

    async with httpx.AsyncClient(timeout=60.0) as client:  # Create new client per request
        try:
            response = await client.post(url, json=payload)
            response.raise_for_status()
            return response.json()
        except httpx.RequestError as e:
            logger.error(f"Request error for {url}: {e}", exc_info=True)
            raise
        except httpx.HTTPStatusError as e:
            logger.error(f"HTTP error {e.response.status_code} for {url}: {e.response.text}")
            raise


def get_db(max_retries=3, retry_delay=1):
    """
    Get database connection with retry logic and timeout configuration.
    Uses Flask's g object to store connection per request.
    
    Args:
        max_retries: Maximum number of connection retry attempts
        retry_delay: Delay in seconds between retries
    
    Returns:
        Database connection object
    """
    if 'db' not in g:
        conn_str = (
            f"DRIVER={{ODBC Driver 18 for SQL Server}};"
            f"SERVER={app.config['AZURE_SQL_SERVER']};"
            f"DATABASE={app.config['AZURE_SQL_DATABASE']};"
            f"UID={app.config['AZURE_SQL_USERNAME']};"
            f"PWD={app.config['AZURE_SQL_PASSWORD']};"
            "Encrypt=yes;TrustServerCertificate=no;"
            "Connection Timeout=30;"  # 30 second connection timeout
        )
        
        # Retry logic for transient connection failures
        last_exception = None
        for attempt in range(max_retries):
            try:
                g.db = pyodbc.connect(conn_str, timeout=30)
                # Set connection timeout for queries
                g.db.timeout = 30
                logger.debug(f"Database connection established (attempt {attempt + 1})")
                break
            except (pyodbc.OperationalError, pyodbc.InterfaceError) as e:
                last_exception = e
                if attempt < max_retries - 1:
                    logger.warning(f"Database connection attempt {attempt + 1} failed: {e}. Retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                else:
                    logger.error(f"Failed to connect to database after {max_retries} attempts: {e}", exc_info=True)
                    raise
            except Exception as e:
                logger.error(f"Unexpected database connection error: {e}", exc_info=True)
                raise
    
    return g.db

# Database cleanup handler
@app.teardown_appcontext
def close_db(error=None):
    """Close database connection at end of request"""
    db = g.pop('db', None)
    if db is not None:
        try:
            db.close()
            logger.debug("Database connection closed")
        except Exception as e:
            logger.warning(f"Error closing database connection: {e}", exc_info=True)

# Database cleanup is now handled by teardown_appcontext in get_db()

def init_db():
    with app.app_context():
        db = get_db()
        cursor = db.cursor()
        
        # Create users table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'users')
        CREATE TABLE users (
            id INT IDENTITY(1,1) PRIMARY KEY,
            username NVARCHAR(255) UNIQUE NOT NULL,
            email NVARCHAR(255) UNIQUE NOT NULL,
            password NVARCHAR(500) NOT NULL,
            firm NVARCHAR(255),
            location NVARCHAR(255),
            lawyer_name NVARCHAR(255),
            state NVARCHAR(50),
            address NVARCHAR(255),
            planning_session NVARCHAR(255),
            other_planning_session NVARCHAR(255),
            discovery_call_link NVARCHAR(255)
        )
        ''')
        
        # Add new columns if they don't exist
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'address')
            ALTER TABLE users ADD address NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'planning_session')
            ALTER TABLE users ADD planning_session NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'other_planning_session')
            ALTER TABLE users ADD other_planning_session NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'discovery_call_link')
            ALTER TABLE users ADD discovery_call_link NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
        
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'selected_tone')
            ALTER TABLE users ADD selected_tone NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'tone_description')
            ALTER TABLE users ADD tone_description NVARCHAR(MAX)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'keywords')
            ALTER TABLE users ADD keywords NVARCHAR(MAX)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'is_blocked')
            ALTER TABLE users ADD is_blocked BIT DEFAULT 0
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'is_admin')
            ALTER TABLE users ADD is_admin BIT DEFAULT 0
            ''')
        except pyodbc.Error:
            pass
        
        # Update password column size to accommodate hashed passwords
        try:
            cursor.execute('''
            ALTER TABLE users ALTER COLUMN password NVARCHAR(500)
            ''')
        except pyodbc.Error:
            pass
        
        # Create tones table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'tones')
        CREATE TABLE tones (
            id INT IDENTITY(1,1) PRIMARY KEY,
            user_id INT NOT NULL,
            name NVARCHAR(255) NOT NULL,
            description NVARCHAR(MAX) NOT NULL,
            CONSTRAINT UQ_user_tone UNIQUE(user_id, name),
            CONSTRAINT FK_user_tone FOREIGN KEY(user_id) REFERENCES users(id)
        )
        ''')
        
        # NOTE: Default users with hardcoded credentials have been removed for security.
        # To create an admin user, use the registration endpoint or create a secure setup script.
        # For production, create admin users through a secure first-run setup process.
        
        # Create password_resets table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='password_resets' AND xtype='U')
        CREATE TABLE password_resets (
            id INTEGER IDENTITY(1,1) PRIMARY KEY,
            email NVARCHAR(255) NOT NULL,
            token NVARCHAR(255) NOT NULL UNIQUE,
            expires DATETIME NOT NULL,
            used INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT GETDATE()
        )
        ''')
        
        # Create feedback table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'feedback')
        CREATE TABLE feedback (
            id INT IDENTITY(1,1) PRIMARY KEY,
            user_id INT,
            feedback_type NVARCHAR(50) NOT NULL,
            priority NVARCHAR(20) NOT NULL,
            subject NVARCHAR(255) NOT NULL,
            message NVARCHAR(MAX) NOT NULL,
            contact_email NVARCHAR(255),
            status NVARCHAR(20) DEFAULT 'pending',
            created_at DATETIME DEFAULT GETDATE(),
            updated_at DATETIME DEFAULT GETDATE(),
            CONSTRAINT FK_feedback_user FOREIGN KEY(user_id) REFERENCES users(id)
        )
        ''')
        
        # Create user_activity table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'user_activity')
        CREATE TABLE user_activity (
            id INT IDENTITY(1,1) PRIMARY KEY,
            user_id INT NOT NULL,
            activity_type NVARCHAR(100) NOT NULL,
            feature_name NVARCHAR(100) NOT NULL,
            api_endpoint NVARCHAR(255),
            request_payload_size INT,
            response_status INT,
            response_size INT,
            processing_time_ms INT,
            success BIT DEFAULT 1,
            error_message NVARCHAR(MAX),
            additional_data NVARCHAR(MAX),
            ip_address NVARCHAR(45),
            user_agent NVARCHAR(500),
            created_at DATETIME DEFAULT GETDATE(),
            CONSTRAINT FK_user_activity_user FOREIGN KEY(user_id) REFERENCES users(id)
        )
        ''')
        
        # Create index for better query performance
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.indexes WHERE name = 'IX_user_activity_user_date')
            CREATE INDEX IX_user_activity_user_date ON user_activity(user_id, created_at)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.indexes WHERE name = 'IX_user_activity_type_date')
            CREATE INDEX IX_user_activity_type_date ON user_activity(activity_type, created_at)
            ''')
        except pyodbc.Error:
            pass
        
        # Create articles table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'articles')
        CREATE TABLE articles (
            id INT IDENTITY(1,1) PRIMARY KEY,
            title NVARCHAR(255) NOT NULL,
            description NVARCHAR(MAX),
            filename NVARCHAR(255) NOT NULL,
            markdown_content NVARCHAR(MAX),
            docx_content VARBINARY(MAX),
            created_at DATETIME2 DEFAULT GETDATE(),
            created_by INT REFERENCES users(id),
            is_active BIT DEFAULT 1,
            status NVARCHAR(50) DEFAULT 'active'
        )
        ''')
        
        # Create index for better query performance
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.indexes WHERE name = 'IX_articles_active_status')
            CREATE INDEX IX_articles_active_status ON articles(is_active, status)
            ''')
        except pyodbc.Error:
            pass
        
        db.commit()

# Initialize database
with app.app_context():
    init_db()

# Add context processor to inject current year into all templates
@app.context_processor
def inject_year():
    return {'now': datetime.now()}

class UserActivityTracker:
    @staticmethod
    def log_activity(user_id, activity_type, feature_name, api_endpoint=None, 
                    request_payload_size=None, response_status=None, response_size=None,
                    processing_time_ms=None, success=True, error_message=None, 
                    additional_data=None, ip_address=None, user_agent=None):
        """Log user activity to the database"""
        try:
            db = get_db()
            cursor = db.cursor()
            
            cursor.execute('''
            INSERT INTO user_activity 
            (user_id, activity_type, feature_name, api_endpoint, request_payload_size, 
             response_status, response_size, processing_time_ms, success, error_message, 
             additional_data, ip_address, user_agent)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (user_id, activity_type, feature_name, api_endpoint, request_payload_size,
                  response_status, response_size, processing_time_ms, success, error_message,
                  additional_data, ip_address, user_agent))
            
            db.commit()
            return True
        except Exception as e:
            logger.error(f"Error logging user activity: {str(e)}", exc_info=True)
            return False
    
    @staticmethod
    def get_user_activity_summary(user_id=None, days=30):
        """Get activity summary for analytics"""
        try:
            db = get_db()
            cursor = db.cursor()
            
            if user_id:
                # Get activity for specific user
                cursor.execute('''
                SELECT 
                    activity_type,
                    feature_name,
                    COUNT(*) as usage_count,
                    AVG(processing_time_ms) as avg_processing_time,
                    SUM(CASE WHEN success = 1 THEN 1 ELSE 0 END) as success_count,
                    SUM(CASE WHEN success = 0 THEN 1 ELSE 0 END) as error_count
                FROM user_activity 
                WHERE user_id = ? AND created_at >= DATEADD(day, -?, GETDATE())
                GROUP BY activity_type, feature_name
                ORDER BY usage_count DESC
                ''', (user_id, days))
            else:
                # Get overall activity summary
                cursor.execute('''
                SELECT 
                    u.username,
                    ua.activity_type,
                    ua.feature_name,
                    COUNT(*) as usage_count,
                    AVG(ua.processing_time_ms) as avg_processing_time,
                    SUM(CASE WHEN ua.success = 1 THEN 1 ELSE 0 END) as success_count,
                    SUM(CASE WHEN ua.success = 0 THEN 1 ELSE 0 END) as error_count
                FROM user_activity ua
                JOIN users u ON ua.user_id = u.id
                WHERE ua.created_at >= DATEADD(day, -?, GETDATE())
                GROUP BY u.username, ua.activity_type, ua.feature_name
                ORDER BY usage_count DESC
                ''', (days,))
            
            return cursor.fetchall()
        except Exception as e:
            logger.error(f"Error getting activity summary: {str(e)}", exc_info=True)
            return []
    
    @staticmethod
    def get_feature_usage_stats(days=30):
        """Get feature usage statistics"""
        try:
            db = get_db()
            cursor = db.cursor()
            
            cursor.execute('''
            SELECT 
                feature_name,
                COUNT(*) as total_usage,
                COUNT(DISTINCT user_id) as unique_users,
                AVG(processing_time_ms) as avg_processing_time,
                SUM(CASE WHEN success = 1 THEN 1 ELSE 0 END) as success_count,
                SUM(CASE WHEN success = 0 THEN 1 ELSE 0 END) as error_count
            FROM user_activity 
            WHERE created_at >= DATEADD(day, -?, GETDATE())
            GROUP BY feature_name
            ORDER BY total_usage DESC
            ''', (days,))
            
            return cursor.fetchall()
        except Exception as e:
            logger.error(f"Error getting feature usage stats: {str(e)}", exc_info=True)
            return []

class UserSession:
    """
    User session management class handling registration, login, and profile operations.
    
    Provides static methods for user authentication, profile management, and session handling.
    All methods include input validation and proper error handling.
    """
    
    @staticmethod
    def register(email: str, password: str, firm: str, location: str, lawyer_name: str, 
                 state: str, address: str = "", planning_session: str = "", 
                 other_planning_session: str = "", discovery_call_link: str = "") -> Tuple[bool, Optional[str]]:
        """
        Register a new user with validation.
        
        Args:
            email: User email address (validated)
            password: User password (must meet strength requirements)
            firm: Law firm name
            location: Firm location
            lawyer_name: Name of the lawyer
            state: State abbreviation
            address: Optional address
            planning_session: Optional planning session name
            other_planning_session: Optional other planning session
            discovery_call_link: Optional discovery call link
            
        Returns:
            Tuple of (success: bool, error_message: Optional[str])
            - (True, None) on success
            - (False, error_message) on failure
        """
        # Validate inputs
        email_valid, email_error = validate_email(email)
        if not email_valid:
            logger.warning(f"Registration failed: {email_error}")
            return False, email_error
        
        password_valid, password_error = validate_password(password)
        if not password_valid:
            logger.warning(f"Registration failed: {password_error}")
            return False, password_error
        
        # Validate and sanitize text inputs
        firm_valid, firm_error = validate_text_input(firm, "Firm name", min_length=1, max_length=200, required=True)
        if not firm_valid:
            logger.warning(f"Registration failed: {firm_error}")
            return False, firm_error
        
        location_valid, location_error = validate_text_input(location, "Location", min_length=1, max_length=200, required=True)
        if not location_valid:
            logger.warning(f"Registration failed: {location_error}")
            return False, location_error
        
        lawyer_name_valid, lawyer_name_error = validate_text_input(lawyer_name, "Lawyer name", min_length=1, max_length=200, required=True)
        if not lawyer_name_valid:
            logger.warning(f"Registration failed: {lawyer_name_error}")
            return False, lawyer_name_error
        
        state_valid, state_error = validate_text_input(state, "State", min_length=2, max_length=50, required=True)
        if not state_valid:
            logger.warning(f"Registration failed: {state_error}")
            return False, state_error
        
        # Sanitize optional fields
        address = sanitize_input(address, max_length=500) if address else ""
        planning_session = sanitize_input(planning_session, max_length=200) if planning_session else ""
        other_planning_session = sanitize_input(other_planning_session, max_length=200) if other_planning_session else ""
        discovery_call_link = sanitize_input(discovery_call_link, max_length=500) if discovery_call_link else ""
        
        db = get_db()
        username = email.split('@')[0].lower()
        try:
            # Hash password before storing
            hashed_password = generate_password_hash(password)
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, address, 
                             planning_session, other_planning_session, discovery_call_link)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (username, email, hashed_password, firm, location, lawyer_name, state, address, 
                 planning_session, other_planning_session, discovery_call_link))
            db.commit()
            logger.info(f"User registered successfully: {email}")
            return True, None
        except pyodbc.IntegrityError:
            logger.warning(f"Registration failed: Email already registered - {email}")
            return False, "Email already registered"
        except Exception as e:
            logger.error(f"Registration error: {str(e)}", exc_info=True)
            return False, "Registration failed. Please try again."

    @staticmethod
    def login(email, password):
        db = get_db()
        username = email.split('@')[0].lower()
        cursor = db.cursor()
        cursor.execute('SELECT * FROM users WHERE username = ?', (username,))
        user = cursor.fetchone()
        
        if not user:
            return False
        
        # Check password - support both hashed (new) and plain text (legacy) passwords
        password_valid = False
        stored_password = user.password
        
        # First, try checking against hashed password (for new passwords)
        # Werkzeug can use pbkdf2:sha256: or scrypt: format
        if stored_password and (stored_password.startswith('pbkdf2:sha256:') or stored_password.startswith('scrypt:')):
            # Password is hashed, use check_password_hash
            password_valid = check_password_hash(stored_password, password)
            # If valid and password was hashed, we're good
        else:
            # Password is stored in plain text (legacy), compare directly
            if stored_password == password:
                password_valid = True
                # Migrate to hashed password for security
                try:
                    hashed_password = generate_password_hash(password)
                    cursor.execute('UPDATE users SET password = ? WHERE id = ?', (hashed_password, user.id))
                    db.commit()
                except Exception as e:
                    # Log error but don't fail login
                    logger.warning(f"Failed to migrate password to hashed format for user {user.id}: {e}", exc_info=True)
        
        if password_valid:
            # Check if user is blocked
            if hasattr(user, 'is_blocked') and user.is_blocked:
                # Log blocked user login attempt
                UserActivityTracker.log_activity(
                    user_id=user.id,
                    activity_type="authentication",
                    feature_name="User Login",
                    api_endpoint="N/A",
                    success=False,
                    error_message="User account is blocked",
                    additional_data=f"Blocked user login attempt from email: {email}"
                )
                return False
            
            # Get user's custom tones
            cursor.execute('SELECT name, description FROM tones WHERE user_id = ?', (user.id,))
            tones = cursor.fetchall()
            
            session['user'] = {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'firm': user.firm,
                'location': user.location,
                'lawyer_name': user.lawyer_name,
                'state': user.state,
                'address': user.address,
                'planning_session': user.planning_session,
                'other_planning_session': user.other_planning_session,
                'discovery_call_link': user.discovery_call_link,
                'is_admin': getattr(user, 'is_admin', False),
                'custom_tones': [{'name': tone.name, 'description': tone.description} for tone in tones]
            }
            
            # Log successful login activity
            UserActivityTracker.log_activity(
                user_id=user.id,
                activity_type="authentication",
                feature_name="User Login",
                api_endpoint="N/A",
                success=True,
                additional_data=f"Login from email: {email}"
            )
            
            return True
        else:
            # Log failed login attempt if user exists
            UserActivityTracker.log_activity(
                user_id=user.id,
                activity_type="authentication",
                feature_name="User Login",
                api_endpoint="N/A",
                success=False,
                error_message="Invalid password",
                additional_data=f"Failed login attempt from email: {email}"
            )
            return False

    @staticmethod
    def update_profile(username, firm, location, lawyer_name, state, address="", planning_session="", other_planning_session="", discovery_call_link="", selected_tone="", tone_description="", keywords=""):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            UPDATE users 
            SET firm = ?, location = ?, lawyer_name = ?, state = ?, 
                address = ?, planning_session = ?, other_planning_session = ?, discovery_call_link = ?,
                selected_tone = ?, tone_description = ?, keywords = ?
            WHERE username = ?
            ''', (firm, location, lawyer_name, state, address, planning_session, 
                 other_planning_session, discovery_call_link, selected_tone, tone_description, keywords, username))
            db.commit()
            
            # Update session if this is the current user
            if 'user' in session and session['user']['username'] == username:
                session['user'].update({
                    'firm': firm,
                    'location': location,
                    'lawyer_name': lawyer_name,
                    'state': state,
                    'address': address,
                    'planning_session': planning_session,
                    'other_planning_session': other_planning_session,
                    'discovery_call_link': discovery_call_link,
                    'selected_tone': selected_tone,
                    'tone_description': tone_description,
                    'keywords': keywords
                })
                session.modified = True
            return True
        except pyodbc.Error:
            return False

    @staticmethod
    def get_current_user():
        return session.get('user')
    
    @staticmethod
    def block_user(user_id, blocked=True):
        """Block or unblock a user account"""
        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute('UPDATE users SET is_blocked = ? WHERE id = ?', (1 if blocked else 0, user_id))
            db.commit()
            
            # Log the blocking action
            UserActivityTracker.log_activity(
                user_id=user_id,
                activity_type="account_management",
                feature_name="User Blocking",
                api_endpoint="N/A",
                success=True,
                additional_data=f"User {'blocked' if blocked else 'unblocked'}"
            )
            
            return True
        except Exception as e:
            logger.error(f"Error blocking/unblocking user: {str(e)}", exc_info=True)
            return False
    
    @staticmethod
    def is_user_blocked(user_id):
        """Check if a user is blocked"""
        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute('SELECT is_blocked FROM users WHERE id = ?', (user_id,))
            result = cursor.fetchone()
            return result and result.is_blocked
        except Exception as e:
            logger.error(f"Error checking user block status: {str(e)}", exc_info=True)
            return False

    @staticmethod
    def add_custom_tone(user_id, tone_name, tone_description):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO tones (user_id, name, description)
            VALUES (?, ?, ?)
            ''', (user_id, tone_name, tone_description))
            db.commit()
            
            # Update session if this is the current user
            if 'user' in session and session['user']['id'] == user_id:
                session['user']['custom_tones'].append({
                    'name': tone_name,
                    'description': tone_description
                })
                session.modified = True
            return True
        except pyodbc.IntegrityError:
            return False
    
    @staticmethod
    def get_custom_tones(user_id):
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT name, description FROM tones WHERE user_id = ?', (user_id,))
        tones = cursor.fetchall()
        return [{'name': tone.name, 'description': tone.description} for tone in tones]

    @staticmethod
    def submit_feedback(user_id, feedback_type, priority, subject, message, contact_email=None):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO feedback (user_id, feedback_type, priority, subject, message, contact_email)
            VALUES (?, ?, ?, ?, ?, ?)
            ''', (user_id, feedback_type, priority, subject, message, contact_email))
            db.commit()
            return True
        except Exception as e:
            logger.error(f"Error submitting feedback: {str(e)}", exc_info=True)
            return False

class Config:
    ARTICLES_DIR = "content/articles"
    GENERATED_DIR = "generated"
    os.makedirs(ARTICLES_DIR, exist_ok=True)
    os.makedirs(GENERATED_DIR, exist_ok=True)

    # Default section markers (can be updated based on client requirements)
    SECTION_MARKERS = {
        'hook': {
            'start': 0,  # First paragraph
            'end': 1     # End of first paragraph
        },
        'summary': {
            'start': 1,  # Second paragraph (2-3 lines ending with "read more...")
            'end': 2     # End of second paragraph
        },
        'disclaimer': {
            'start': -1,  # Last paragraph (disclaimer)
            'end': None   # End of content
        }
    }

class AzureServices:
    """
    Service class for interacting with Azure OpenAI and content generation services.
    
    Handles content rewriting, validation, and section preservation for blog articles.
    """
    
    def __init__(self):
        self.text_client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_KEY"),
            api_version="2024-02-15-preview",
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )
        
        self.conversations = {}

    def _extract_sections(self, content):
        """Extract and preserve specific sections from the content."""
        try:
            paragraphs = content.split('\n\n')
            preserved_sections = {}
            
            logger.debug(f"Extracting sections - Total paragraphs found: {len(paragraphs)}")
            
            # Extract Hook (first paragraph)
            hook_start = Config.SECTION_MARKERS['hook']['start']
            hook_end = Config.SECTION_MARKERS['hook']['end']
            if hook_start < len(paragraphs):
                preserved_sections['hook'] = paragraphs[hook_start:hook_end][0] if hook_end - hook_start == 1 else '\n\n'.join(paragraphs[hook_start:hook_end])
                logger.debug(f"Extracted Hook (paragraph 1): {preserved_sections['hook'][:100]}...")
            else:
                logger.warning("Hook section not found in content")
                preserved_sections['hook'] = ""
            
            # Extract Summary (second paragraph - should be 2-3 lines ending with "read more...")
            summary_start = Config.SECTION_MARKERS['summary']['start']
            summary_end = Config.SECTION_MARKERS['summary']['end']
            if summary_start < len(paragraphs):
                preserved_sections['summary'] = paragraphs[summary_start:summary_end][0] if summary_end - summary_start == 1 else '\n\n'.join(paragraphs[summary_start:summary_end])
                logger.debug(f"Extracted Summary (paragraph 2): {preserved_sections['summary'][:100]}...")
            else:
                logger.warning("Summary section not found in content")
                preserved_sections['summary'] = ""
            

            
            # Extract Disclaimer (last paragraph)
            disclaimer_start = Config.SECTION_MARKERS['disclaimer']['start']
            if len(paragraphs) > 0:
                preserved_sections['disclaimer'] = paragraphs[disclaimer_start]
                logger.debug(f"Extracted Disclaimer (paragraph {len(paragraphs)}): {preserved_sections['disclaimer'][:100]}...")
            else:
                logger.warning("Disclaimer section not found in content")
                preserved_sections['disclaimer'] = ""
            
            logger.debug("Section extraction complete")
            
            return preserved_sections
        except Exception as e:
            logger.error(f"Error extracting sections: {str(e)}", exc_info=True)
            return {'hook': "", 'summary': "", 'disclaimer': ""}

    def _reconstruct_content(self, new_content, preserved_sections):
        """Preserve specific sections exactly as they are in the original content."""
        try:
            paragraphs = new_content.split('\n\n')
            
            logger.debug(f"Reconstructing content - Total paragraphs in new content: {len(paragraphs)}")
            
            # Remove any existing disclaimer or service description from the middle of the content
            # to prevent duplication, but preserve CTA content
            cleaned_paragraphs = []
            for i, para in enumerate(paragraphs):
                # Skip paragraphs that are exact matches of disclaimers or service descriptions
                # But preserve paragraphs that contain CTA phrases
                is_duplicate_section = False
                if preserved_sections['disclaimer'] and preserved_sections['disclaimer'].strip() == para.strip():
                    is_duplicate_section = True
                    logger.debug(f"Removed duplicate disclaimer from position {i+1}")

                
                if not is_duplicate_section:
                    cleaned_paragraphs.append(para)
            
            paragraphs = cleaned_paragraphs
            logger.debug(f"Paragraphs after cleaning: {len(paragraphs)}")
            
            # Place preserved sections in their correct positions
            # Hook (first paragraph)
            if preserved_sections['hook'] and len(paragraphs) > 0:
                paragraphs[0] = preserved_sections['hook']
                logger.debug("Preserved Hook section at position 1")
            
            # Summary (second paragraph - 2-3 lines ending with "read more...")
            if preserved_sections['summary'] and len(paragraphs) > 1:
                paragraphs[1] = preserved_sections['summary']
                logger.debug("Preserved Summary section at position 2")
            elif preserved_sections['summary'] and len(paragraphs) <= 1:
                # If there aren't enough paragraphs, add the summary as second paragraph
                if len(paragraphs) == 0:
                    paragraphs.append("")  # Add empty first paragraph if needed
                paragraphs.append(preserved_sections['summary'])
                logger.debug("Added Summary section at position 2")
            

            
            # Disclaimer (last paragraph)
            if preserved_sections['disclaimer']:
                # Always add disclaimer at the very end
                paragraphs.append(preserved_sections['disclaimer'])
                logger.debug("Added Disclaimer section at the very end")
            
            final_content = '\n\n'.join(paragraphs)
            
            # Verify sections are preserved exactly
            if preserved_sections['hook'] and preserved_sections['hook'] not in final_content:
                logger.warning("Hook section not preserved exactly")
            if preserved_sections['summary'] and preserved_sections['summary'] not in final_content:
                logger.warning("Summary section not preserved exactly")

            if preserved_sections['disclaimer'] and preserved_sections['disclaimer'] not in final_content:
                logger.warning("Disclaimer section not preserved exactly")
            
            logger.debug("Content reconstruction complete")
            
            return final_content
        except Exception as e:
            logger.error(f"Error preserving sections: {str(e)}", exc_info=True)
            return new_content

    def _validate_with_gpt(self, original_text, new_content, components):
        """Validate article components using GPT for better semantic understanding."""
        validation_prompt = f"""
            You are an expert content validator. Analyze these two articles and provide a detailed validation.
            You MUST respond with a valid JSON object following this EXACT structure, with no additional text:
            
            {{
                "components": {{
                    "keywords": {{
                        "found": true/false,
                        "occurrences": number,
                        "variations": ["variation1", "variation2"],
                        "in_first_150": true/false
                    }},
                    "firm_info": {{
                        "found": true/false,
                        "name": true/false,
                        "location": true/false
                    }},
                    "lawyer_info": {{
                        "found": true/false,
                        "name": true/false,
                        "location": true/false
                    }},
                    "planning_session": {{
                        "found": true/false,
                        "name": true/false,
                        "references": number
                    }},
                    "discovery_call": {{
                        "found": true/false,
                        "link": true/false,
                        "references": number
                    }}
                }},
                "preserved_sections": {{
                    "hook": true/false,
                    "summary": true/false,
                    "disclaimer": true/false
                }},
                "change_analysis": {{
                    "percentage": number,
                    "significant_changes": true/false,
                    "maintained_essence": true/false
                }},
                "warnings": ["warning1", "warning2"],
                "missing_components": ["component1", "component2"]
            }}

            Analyze the following content:

            Required components to check:
            - Keywords: {components['keywords']}
            - Firm: {components['firm_name']} in {components['location']}
            - Lawyer: {components['lawyer_name']} in {components['city']}, {components['state']}
            - Planning Session: {components['planning_session_name']}
            - Discovery Call: {components['discovery_call_link']}

            Original Article:
            {original_text}

            New Article:
            {new_content}

            Remember to respond with ONLY the JSON object, no additional text or explanation.
        """

        try:
            response = self.text_client.chat.completions.create(
                model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                messages=[
                    {"role": "system", "content": "You are a JSON-only response validator. Always respond with valid JSON matching the exact structure provided."},
                    {"role": "user", "content": validation_prompt}
                ],
                temperature=0.1,  # Lower temperature for more consistent JSON output
                response_format={ "type": "json_object" }  # Force JSON response
            )

            # Get the response content and ensure it's valid JSON
            response_content = response.choices[0].message.content.strip()
            
            # Try to parse the JSON response
            try:
                validation_results = json.loads(response_content)
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing JSON response: {str(e)}", exc_info=True)
                logger.debug(f"Raw response: {response_content}")
                raise

            # Validate the structure of the response
            required_keys = ['components', 'preserved_sections', 'change_analysis', 'warnings', 'missing_components']
            if not all(key in validation_results for key in required_keys):
                logger.error("Invalid response structure. Missing required keys.")
                raise ValueError("Invalid response structure")

            # Log validation results
            logger.debug("GPT Article Validation Results:")
            for component, details in validation_results['components'].items():
                status = '✓' if details.get('found', False) else '✗'
                logger.debug(f"Component {component}: {status}")
                if component == 'keywords' and details.get('variations'):
                    logger.debug(f"  Variations found: {', '.join(details['variations'])}")
                if 'occurrences' in details:
                    logger.debug(f"  Occurrences: {details['occurrences']}")

            for section, preserved in validation_results['preserved_sections'].items():
                logger.debug(f"Preserved section {section}: {'✓' if preserved else '✗'}")

            logger.debug(f"Change Analysis - Percentage: {validation_results['change_analysis']['percentage']:.1f}%")
            logger.debug(f"Significant Changes: {'✓' if validation_results['change_analysis']['significant_changes'] else '✗'}")
            logger.debug(f"Maintained Essence: {'✓' if validation_results['change_analysis']['maintained_essence'] else '✗'}")

            if validation_results['warnings']:
                for warning in validation_results['warnings']:
                    logger.warning(f"Validation warning: {warning}")

            if validation_results['missing_components']:
                for component in validation_results['missing_components']:
                    logger.warning(f"Missing component: {component}")

            return validation_results

        except Exception as e:
            logger.error(f"Error in GPT validation: {str(e)}", exc_info=True)
            logger.warning("Unable to validate article components. Please check the generated content manually.")
            return None

    def rewrite_content(self, original_text, tone, tone_description, keywords, firm_name, location, lawyer_name, city, state, discovery_call_link, planning_session_name="15-minute discovery call"):
        try:
            # Extract sections to preserve
            logger.debug("Extracting sections to preserve...")
            preserved_sections = self._extract_sections(original_text)
            
            # CRITICAL: DO NOT MODIFY THESE SECTIONS {preserved_sections}:
            # 1. The first paragraph (Hook) - which is this
            # 2. The fourth paragraph (Plug) - Keep it exactly as is
            # 3. The last paragraph (Disclaimer) - Keep it exactly as is
            
            # CRITICAL: DO NOT REPEAT THESE SECTIONS:
            # 1. The hook paragraph should appear ONLY ONCE at the beginning
            # 2. The plug paragraph should appear ONLY ONCE in its original position
            # 3. The disclaimer paragraph should appear ONLY ONCE at the end
            # 4. DO NOT include these preserved sections anywhere else in the article
            # 5. DO NOT create new paragraphs that repeat the same content as the hook, plug, or disclaimer

            # Add explicit instructions about preserving sections
            # CRITICAL: DO NOT REPEAT THESE SECTIONS:
            # 1. The first paragraph should appear ONLY ONCE at the beginning
            # 2. The second paragraph should appear ONLY ONCE in its original position
            # 3. The disclaimer paragraph should appear ONLY ONCE at the end
            # 4. DO NOT include these preserved sections anywhere else in the article
            # 5. DO NOT create new paragraphs that repeat the same content as the hook, plug, or disclaimer


            system_prompt = f"""
                You are a legal blog post rewriter. Generate ONLY the main article content (heading + body + CTA) with at least 40% changes from the original.
                
                TEMPLATE STRUCTURE (DO NOT INCLUDE THESE SECTIONS - THEY WILL BE ADDED AUTOMATICALLY):
                - Hook: {preserved_sections['hook']}
                - Summary: {preserved_sections['summary']}
                - Date: Will be added automatically
                - Disclaimer: {preserved_sections['disclaimer']}
                
                YOUR TASK: Generate ONLY the main content that goes in the template:
                1. Main heading (starts with "# ")
                2. Article body with subheadings (## )
                3. Call-to-action (CTA)
                
                The content will be inserted into a template that already includes:
                - Summary section
                - Date section  
                - Disclaimer section
                
                CRITICAL REQUIREMENTS:
                1. Start your content with the main heading: "# [Title]"
                2. Use proper markdown formatting throughout
                3. Include 3-4 subheadings with "## "
                4. End with a CTA paragraph
                5. DO NOT include hook, summary, date, or disclaimer
                6. DO NOT include any preview text or introductory paragraphs
                7. DO NOT include any dates
                8. Ensure at least 40% changes from original
                
                FORMATTING REQUIREMENTS:
                - Main heading: "# [Title]"
                - Subheadings: "## [Subheading]"
                - Bold text: **text**
                - Italic text: *text*
                - Bullet points: - or *
                - Proper line breaks between paragraphs
                
                SEO REQUIREMENTS:
                - Include keywords: {keywords}
                - Mention firm: {firm_name} in {location}
                - Mention lawyer: {lawyer_name} in {city}, {state}
                - Include planning session: {planning_session_name}
                - Include discovery call link: {discovery_call_link}
                
                TONE: {tone} - {tone_description}
                
                CTA REQUIREMENTS:
                - Use "15-minute discovery call" (lowercase) as clickable text
                - Format as markdown link: [15-minute discovery call]({discovery_call_link})
                - Include this link in the CTA paragraph
                - End content immediately after CTA
                
                LINK TEXT REQUIREMENTS:
                - Use descriptive link text that explains what the link does
                - DO NOT use generic phrases like "click here", "read more", "learn more"
                - Use specific, action-oriented text like "schedule your consultation", "book your session", "get started today"
                
                Generate ONLY the main content (heading + body + CTA). The system will add hook, summary, date, and disclaimer automatically.
            """
            
            logger.debug("Generating rewritten content...")
            response = self.text_client.chat.completions.create(
                model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": original_text}
                ],
                temperature=0.7,
            )
            
            # Get the rewritten content
            rewritten_content = response.choices[0].message.content
            
            # Generate summary (2-3 sentences ending with "Read more...")
            logger.debug("Generating summary...")
            summary = self._generate_summary(rewritten_content, preserved_sections['hook'])
            
            # Final assembly with template
            logger.debug("Assembling final article with template...")
            final_content = self._assemble_final_article(preserved_sections['hook'], summary, rewritten_content, preserved_sections['disclaimer'], firm_name, discovery_call_link)
            
            # Validate the generated content using GPT
            components = {
                'keywords': keywords,
                'firm_name': firm_name,
                'location': location,
                'lawyer_name': lawyer_name,
                'city': city,
                'state': state,
                'planning_session_name': planning_session_name,
                'discovery_call_link': discovery_call_link
            }
            
            validation_results = self._validate_with_gpt(original_text, final_content, components)
            
            if validation_results is None:
                logger.warning("Article validation failed. Please review the content manually.")
            
            logger.info("Article generation complete!")
            return final_content
            
        except Exception as e:
            logger.error(f"Error in rewrite_content: {str(e)}", exc_info=True)
            return original_text

    def edit_content(self, session_id, user_message, current_content=None):
        if session_id not in self.conversations:
            self.conversations[session_id] = [
                {"role": "system", "content": """
                    You are a legal blog post editor. When the user requests changes:
                    1. The first paragraph should not be repeating and be the same as the original
                    2. The second paragraph should not be extending and be the same as the original
                    3. The last paragraph (disclaimer) should not be repeating and be the same as the original
                    1. Make ONLY the requested changes
                    2. Return the COMPLETE updated blog (not just updated part) in markdown format
                    3. Don't include any commentary or explanations
                    4. Preserve all formatting and structure
                    5. Don't repeat any paragraph meaning no same paragraph should be present.
                """}
            ]
        
        if current_content:
            self.conversations[session_id].append(
                {"role": "assistant", "content": current_content}
            )
        
        self.conversations[session_id].append(
            {"role": "user", "content": user_message}
        )
        
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=self.conversations[session_id],
            temperature=0.5
        )
        
        ai_response = response.choices[0].message.content
        self.conversations[session_id].append(
            {"role": "assistant", "content": ai_response}
        )
        
        return ai_response
    
    def _validate_and_cleanup_structure(self, content, preserved_sections):
        """Validate and cleanup the article structure to ensure it follows the exact format."""
        try:
            logger.debug("Validating and cleaning article structure")
            
            paragraphs = content.split('\n\n')
            cleaned_paragraphs = []
            
            # Step 1: Ensure hook is first
            if preserved_sections['hook']:
                cleaned_paragraphs.append(preserved_sections['hook'])
                logger.debug("Hook placed at position 1")
            else:
                logger.warning("No hook found")
            
            # Step 2: Add line break after hook
            cleaned_paragraphs.append("")
            logger.debug("Added line break after hook")
            
            # Step 3: Ensure summary is second (2-3 lines)
            if preserved_sections['summary']:
                cleaned_paragraphs.append(preserved_sections['summary'])
                logger.debug("Summary placed at position 3")
            else:
                logger.warning("No summary found")
            
            # Step 4: Add line break after summary
            cleaned_paragraphs.append("")
            logger.debug("Added line break after summary")
            
            # Step 5: Add date
            current_date = datetime.now().strftime("%B %d, %Y")
            cleaned_paragraphs.append(f"**Date: {current_date}**")
            logger.debug("Added date")
            
            # Step 6: Add line break after date
            cleaned_paragraphs.append("")
            logger.debug("Added line break after date")
            
            # Step 7: Find and organize content (heading, article content, CTA)
            cta_found = False
            heading_found = False
            content_paragraphs = []
            
            for i, para in enumerate(paragraphs):
                # Skip if this is the hook or summary (already handled)
                if para.strip() == preserved_sections.get('hook', '').strip():
                    continue
                if para.strip() == preserved_sections.get('summary', '').strip():
                    continue
                
                # Skip any paragraphs that look like they might be hook or summary
                if any(phrase in para.lower() for phrase in ['weekly blog preview', 'read more', 'summary', 'featured article section']):
                    continue
                
                # Skip any paragraphs that are just the hook or summary content
                if para.strip() == preserved_sections.get('hook', '').strip():
                    continue
                if para.strip() == preserved_sections.get('summary', '').strip():
                    continue
                
                # Skip any date paragraphs (we'll add our own date)
                if any(phrase in para.lower() for phrase in ['date:', '**date:', '2025.', '2024.', '2023.']):
                    continue
                
                # Check if this paragraph contains CTA
                cta_phrases = ["Click here to schedule", "Book your Discovery Call", "Schedule your complimentary"]
                if any(phrase in para for phrase in cta_phrases):
                    cta_found = True
                    content_paragraphs.append(para)  # Include the CTA paragraph
                    logger.debug(f"Found CTA at paragraph {i+1}")
                    break
                
                # Check if this is a heading (starts with #)
                if para.strip().startswith('#'):
                    if not heading_found:
                        heading_found = True
                        logger.debug(f"Found heading at paragraph {i+1}: {para.strip()[:50]}...")
                    content_paragraphs.append(para)
                else:
                    # Add all other content paragraphs
                    content_paragraphs.append(para)
            
            # Add all content paragraphs (including heading and main content)
            cleaned_paragraphs.extend(content_paragraphs)
            
            if not heading_found:
                logger.warning("No heading found in content")
            if not cta_found:
                logger.warning("No CTA found in content")
            
            # Step 8: Add disclaimer at the end (only if it exists)
            if preserved_sections['disclaimer']:
                cleaned_paragraphs.append(preserved_sections['disclaimer'])
                logger.debug("Disclaimer placed at the end")
            else:
                logger.warning("No disclaimer found")
            
            # Step 9: DROP EVERYTHING AFTER THE DISCLAIMER
            # The disclaimer should be the final paragraph - no additional content after it
            logger.debug("Dropped all content after disclaimer - disclaimer is the final paragraph")
            
            # Clean up any empty paragraphs and ensure proper spacing
            final_paragraphs = []
            for para in cleaned_paragraphs:
                if para.strip():  # Only add non-empty paragraphs
                    final_paragraphs.append(para)
                elif final_paragraphs and final_paragraphs[-1].strip():  # Add empty line only if previous paragraph wasn't empty
                    final_paragraphs.append("")
            
            final_content = '\n\n'.join(final_paragraphs)
            
            logger.debug(f"Structure validation complete. Total paragraphs: {len(final_paragraphs)}")
            logger.debug("Final structure: Hook → Line Break → Summary → Line Break → Date → Line Break → Heading → Content → CTA → Disclaimer")
            
            return final_content
            
        except Exception as e:
            logger.error(f"Error in structure validation: {str(e)}", exc_info=True)
            return content

    def _generate_summary(self, article_content, hook):
        """Generate a 2-3 sentence summary ending with 'Read more...'"""
        try:
            logger.debug("Generating summary")
            
            summary_prompt = f"""
                Generate a 2-3 sentence summary of this article that:
                1. Is engaging and captures the main point
                2. Ends with "Read more..."
                3. Is different from the hook paragraph
                4. Provides a brief overview of what the article covers
                5. Is written in a professional tone
                
                Hook paragraph: {hook}
                
                Article content: {article_content[:1000]}
                
                Generate ONLY the summary (2-3 sentences ending with "Read more..."). Do not include any other text.
            """
            
            response = self.text_client.chat.completions.create(
                model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                messages=[
                    {"role": "system", "content": "You are a summary generation expert. Always return concise, engaging summaries."},
                    {"role": "user", "content": summary_prompt}
                ],
                temperature=0.7,
            )
            
            summary = response.choices[0].message.content.strip()
            
            # Ensure it ends with "Read more..."
            if not summary.endswith("Read more..."):
                summary = summary.rstrip('.') + ". Read more..."
            
            logger.debug(f"Generated summary: {summary[:100]}...")
            logger.debug("Summary generation complete")
            
            return summary
            
        except Exception as e:
            logger.error(f"Error generating summary: {str(e)}", exc_info=True)
            # Fallback summary
            return "This article explores important considerations for protecting your family's future. Read more..."

    def _clean_article_content(self, article_content):
        """Clean and filter article content to remove unwanted elements"""
        try:
            logger.debug("Cleaning article content")
            
            if not article_content:
                return ""
            
            # Split into lines for processing
            lines = article_content.split('\n')
            cleaned_lines = []
            
            # Filter out unwanted content
            skip_until_heading = True
            
            for line in lines:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # Skip preview/summary text
                if any(phrase in line.lower() for phrase in [
                    'weekly blog preview',
                    'summary you can use',
                    'featured article section',
                    'email newsletter',
                    'preview/summary'
                ]):
                    logger.debug(f"Skipped preview text: {line[:50]}...")
                    continue
                
                # Skip any date lines (we'll add our own)
                if any(phrase in line.lower() for phrase in [
                    'date:', '**date:', '2025.', '2024.', '2023.'
                ]):
                    logger.debug(f"Skipped date line: {line[:50]}...")
                    continue
                
                # Skip hook and summary content if they appear in the middle
                if 'read more...' in line.lower() and len(line) < 200:
                    logger.debug(f"Skipped duplicate summary: {line[:50]}...")
                    continue
                
                # Look for the main heading (should start with #)
                if line.startswith('#'):
                    skip_until_heading = False
                    logger.debug(f"Found heading: {line[:50]}...")
                elif skip_until_heading and (
                    # Look for title-like content (longer lines that seem like titles)
                    (len(line) > 20 and len(line) < 100 and 
                     any(word in line.lower() for word in ['mail carrier', 'pet', 'trust', 'planning', 'legacy', 'estate', 'understanding', 'how', 'why', 'what', 'when', 'where', 'guide', 'complete', 'essential', 'important', 'protect', 'secure', 'plan', 'future', 'family', 'legal', 'law', 'attorney', 'lawyer']) and
                     not any(phrase in line.lower() for phrase in ['read more', 'summary', 'preview', 'date:', '**date:'])
                    )
                ):
                    skip_until_heading = False
                    # Ensure proper heading format
                    line = f"# {line}"
                    logger.debug(f"Found heading: {line[:50]}...")
                
                # Only add lines after we've found the heading
                if not skip_until_heading:
                    cleaned_lines.append(line)
            
            # Join lines back together with proper spacing
            cleaned_content = '\n\n'.join(cleaned_lines)
            
            logger.debug("Content cleaned successfully")
            logger.debug("Content cleaning complete")
            
            return cleaned_content
            
        except Exception as e:
            logger.error(f"Error cleaning article content: {str(e)}", exc_info=True)
            return article_content

    def _assemble_final_article(self, hook, summary, article_content, disclaimer, firm_name="", discovery_call_link=""):
        """Assemble the final article using external template file"""
        try:
            logger.debug("Assembling final article with external template")
            
            # Get current date
            current_date = datetime.now().strftime("%B %d, %Y")
            
            # Clean and filter the article content to remove any unwanted elements
            cleaned_content = self._clean_article_content(article_content)
            
            # Read the external template file
            template_path = os.path.join(Config.ARTICLES_DIR, 'markdown', 'template.md')
            try:
                with open(template_path, 'r', encoding='utf-8') as f:
                    template_content = f.read()
                logger.debug("External template loaded successfully")
            except FileNotFoundError:
                logger.warning("Template file not found, using fallback template")
                template_content = """Weekly blog preview/summary you can use on your main blog page and Featured Article section of your email newsletter:

{summary of the article}

**Date: {current_date}**
{newly generated Title with proper markdown formatting}

{newly generated content with proper markdown formatting}

*###### This article is a service of [ name ], a Personal Family Lawyer® Firm. We don't just draft documents; we ensure you make informed and empowered decisions about life and death, for yourself and the people you love. That's why we offer a Life & Legacy Planning® Session, during which you will get more financially organized than you've ever been before and make all the best choices for the people you love. You can begin by calling our office today to schedule a Life & Legacy Planning Session.*

The content is sourced from Personal Family Lawyer® for use by Personal Family Lawyer firms, a source believed to be providing accurate information. This material was created for educational and informational purposes only and is not intended as ERISA, tax, legal, or investment advice. If you are seeking legal advice specific to your needs, such advice services must be obtained on your own separate from this educational material."""
            
            # Replace placeholders in the template
            final_content = template_content.replace('{summary of the article}', summary.strip() if summary else '')
            final_content = final_content.replace('{current_date}', current_date)
            final_content = final_content.replace('{newly generated Title with proper markdown formatting}', '')
            final_content = final_content.replace('{newly generated content with proper markdown formatting}', cleaned_content if cleaned_content else '')
            final_content = final_content.replace('{firm_name}', firm_name if firm_name else '[Firm Name]')
            final_content = final_content.replace('{discovery_call_link}', discovery_call_link if discovery_call_link else '#')
            
            # Remove the preview text line if no hook is provided
            if not hook or not hook.strip():
                lines = final_content.split('\n')
                # Remove the first line (preview text) and any empty lines after it
                if lines and 'Weekly blog preview/summary' in lines[0]:
                    lines = lines[1:]  # Remove first line
                    # Remove empty lines after the removed line
                    while lines and not lines[0].strip():
                        lines = lines[1:]
                final_content = '\n'.join(lines)
            
            logger.debug("Article assembled using external template")
            logger.debug("External template assembly complete")
            
            return final_content
            
        except Exception as e:
            logger.error(f"Error assembling final article: {str(e)}", exc_info=True)
            return article_content

    def _format_markdown(self, content):
        """Apply final markdown formatting to ensure proper structure and formatting."""
        try:
            logger.debug("Applying final markdown formatting")
            
            # Extract the preserved sections first
            preserved_sections = self._extract_sections(content)
            
            formatting_prompt = f"""
                You are a markdown formatting expert. Format the following article content to ensure proper markdown structure and formatting.
                
                CRITICAL: DO NOT MODIFY THESE PRESERVED SECTIONS:
                - Hook: "{preserved_sections.get('hook', '')}"
                - Summary: "{preserved_sections.get('summary', '')}"
                - Disclaimer: "{preserved_sections.get('disclaimer', '')}"
                
                REQUIREMENTS:
                1. DO NOT change the hook, summary, or disclaimer content - keep them exactly as they are
                2. DO NOT move, modify, or reorder the hook, summary, or disclaimer sections
                3. Ensure all headings start with # (main heading), ## (subheadings), or ### (sub-subheadings)
                4. Ensure proper line breaks between sections (use double line breaks)
                5. Format bold text with **text**
                6. Format italic text with *text*
                7. Ensure bullet points use - or * with proper spacing
                8. Ensure the date is properly formatted as **Date: Month DD, YYYY**
                9. Maintain the exact structure: Hook → Line Break → Summary → Line Break → Date → Line Break → Heading → Content → CTA → Line Break → Disclaimer
                10. Do not change the content, only improve the formatting
                11. Ensure proper spacing between paragraphs
                12. Make sure the heading (title) is properly formatted with # at the beginning
                13. Ensure the date line is properly formatted with **Date: Month DD, YYYY**
                14. Add proper line breaks between all sections
                
                CONTENT TO FORMAT:
                {content}
                
                Return ONLY the formatted content with proper markdown formatting. Do not add any explanations or additional text.
            """
            
            response = self.text_client.chat.completions.create(
                model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                messages=[
                    {"role": "system", "content": "You are a markdown formatting expert. Always return properly formatted markdown content."},
                    {"role": "user", "content": formatting_prompt}
                ],
                temperature=0.1,  # Low temperature for consistent formatting
            )
            
            formatted_content = response.choices[0].message.content.strip()
            
            # Validate that preserved sections are still in their correct positions
            logger.debug("Validating preserved sections after formatting")
            final_content = self._validate_and_cleanup_structure(formatted_content, preserved_sections)
            
            logger.debug("Markdown formatting applied successfully")
            logger.debug("Markdown formatting complete")
            
            return final_content
            
        except Exception as e:
            logger.error(f"Error in markdown formatting: {str(e)}", exc_info=True)
            logger.debug("Returning original content without markdown formatting")
            return content

class ImageGenerator:
    def __init__(self):
        self.image_client = AzureOpenAI(
            api_key=os.getenv("AZURE_DALLE_KEY"),
            api_version="2024-02-01",
            azure_endpoint=os.getenv("AZURE_DALLE_ENDPOINT")
        )
        self.text_client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_KEY"),
            api_version="2024-02-15-preview",
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )

    def generate_image(self, text_prompt):
        try:
            safe_prompt = self._get_safe_image_prompt(text_prompt)
            
            response = self.image_client.images.generate(
                model=os.getenv("AZURE_DALLE_DEPLOYMENT"),
                prompt=safe_prompt,
                size="1024x1024",
                quality="standard",
                n=1,
            )
            image_url = response.data[0].url
            os.makedirs(os.path.join(app.static_folder, 'generated'), exist_ok=True)
            
            timestamp = int(time.time())
            image_filename = f"image_{timestamp}.png"
            image_path = os.path.join(app.static_folder, 'generated', image_filename)
            
            response = requests.get(image_url)
            with open(image_path, 'wb') as f:
                f.write(response.content)
            
            return image_filename
            
        except Exception as e:
            logger.error(f"Image generation failed: {e}", exc_info=True)
            return None
        
    def _get_safe_image_prompt(self, text_prompt):
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=[
                {"role": "system", "content": """
                    You are a creative prompt engineer for legal blog images. Create safe and professional image prompts that:
                    1. Are directly relevant to the blog content
                    2. Be 'unique to the blog's content', not generic or reusable for any legal article
                    3. Reflect the main topic, themes, or message of the blog post
                    4. Focus on modern, visually appealing representations
                    5. Must pass Azure content filters
                    6. Avoids sensitive content
                    The prompt should be detailed and specific, including:
                        - Main subject
                        - Style description
                        - Color palette
                        - Composition notes
                        - Mood/tone
                    - Is based on this blog content:
                """},
                {"role": "user", "content": text_prompt[:1000]}
            ],
            temperature=1
        )
        return response.choices[0].message.content

class FileManager:
    @staticmethod
    def list_articles():
        """
        List all articles (from database and file system)
        Returns:
            List of article filenames
        """
        # Get articles from database
        db_articles = set()
        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute("""
                SELECT filename
                FROM articles 
                WHERE is_active = 1 AND status = 'active'
            """)
            for row in cursor.fetchall():
                filename = row[0] if not hasattr(row, 'keys') else row['filename']
                db_articles.add(filename)
        except Exception as e:
            logger.error(f"Error getting articles from database: {str(e)}", exc_info=True)
        
        # Get articles from file system
        fs_articles = set()
        try:
            docx_dir = os.path.join(Config.ARTICLES_DIR, 'docx')
            fs_articles = set([f for f in os.listdir(docx_dir) if f.endswith('.docx')])
        except Exception as e:
            logger.error(f"Error getting articles from file system: {str(e)}", exc_info=True)
        
        # Combine both sources
        all_articles = list(db_articles.union(fs_articles))
        logger.debug(f"list_articles() found {len(all_articles)} total articles")
        logger.debug(f"Database articles: {list(db_articles)}")
        logger.debug(f"File system articles: {list(fs_articles)}")
        
        return all_articles
    
    @staticmethod
    def get_article_metadata():
        """
        Read article metadata from database
        Returns:
            Dictionary of article metadata
        """
        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute("""
                SELECT id, title, description, filename, status
                FROM articles 
                WHERE is_active = 1 AND status = 'active'
                ORDER BY created_at DESC
            """)
            articles = cursor.fetchall()
            
            logger.debug(f"Found {len(articles)} articles in database")
            if articles:
                logger.debug(f"First article type: {type(articles[0])}")
            
            result = {}
            for i, article in enumerate(articles):
                try:
                    # Handle pyodbc.Row, dict, and tuple formats
                    if hasattr(article, 'keys'):  # dict-like (pyodbc.Row or dict)
                        filename = article['filename']
                        result[filename] = {
                            'id': article['id'],
                            'title': article['title'],
                            'description': article['description'],
                            'status': article['status']
                        }
                    else:
                        # Handle tuple format (id, title, description, filename, status)
                        filename = article[3]  # filename is 4th column
                        result[filename] = {
                            'id': article[0],      # id
                            'title': article[1],   # title
                            'description': article[2],  # description
                            'status': article[4]   # status
                        }
                    logger.debug(f"Successfully processed article {i}: {filename}")
                except Exception as e:
                    logger.warning(f"Error processing article {i}: {str(e)}", exc_info=True)
                    continue
            
            logger.debug(f"Final result has {len(result)} articles")
            
            # Also load existing articles from metadata.json to preserve descriptions
            metadata_path = os.path.join(Config.ARTICLES_DIR, 'metadata.json')
            try:
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    metadata = json.loads(content)
                    for article in metadata['articles']:
                        filename = article['filename']
                        if filename not in result:  # Only add if not already in database
                            result[filename] = article
                        else:
                            # Merge descriptions from metadata.json if database description is empty
                            if not result[filename].get('description') and article.get('description'):
                                result[filename]['description'] = article['description']
                logger.debug(f"After merging metadata.json, final result has {len(result)} articles")
                logger.debug(f"Final article list: {list(result.keys())}")
            except (FileNotFoundError, json.JSONDecodeError, KeyError) as e:
                logger.warning(f"Could not load metadata.json: {str(e)}", exc_info=True)
            
            return result
        except Exception as e:
            logger.error(f"Error reading article metadata from database: {str(e)}", exc_info=True)
            # Fallback to metadata.json for backward compatibility
            metadata_path = os.path.join(Config.ARTICLES_DIR, 'metadata.json')
            try:
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    metadata = json.loads(content)
                    result = {article['filename']: article for article in metadata['articles']}
                    return result
            except (FileNotFoundError, json.JSONDecodeError, KeyError) as e:
                logger.error(f"Error reading metadata.json fallback: {str(e)}", exc_info=True)
                return {}
    
    @staticmethod
    def read_docx(filename):
        """
        Read content from a DOCX file (database first, then file system fallback)
        Args:
            filename: Name of the DOCX file (may be URL-encoded)
        Returns:
            Extracted text content
        """
        # URL-decode the filename first
        decoded_filename = unquote(filename)
        
        try:
            # Try to read from database first
            db = get_db()
            cursor = db.cursor()
            cursor.execute("""
                SELECT markdown_content, docx_content
                FROM articles 
                WHERE filename = ? AND is_active = 1 AND status = 'active'
            """, (decoded_filename,))
            article = cursor.fetchone()
            
            if article and (article[0] if not hasattr(article, 'keys') else article['markdown_content']):
                # Convert markdown to plain text for consistency
                import re
                # Get markdown content
                markdown_content = article[0] if not hasattr(article, 'keys') else article['markdown_content']
                # Remove markdown formatting
                text = re.sub(r'#+\s*', '', markdown_content)  # Remove headers
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Remove bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)  # Remove italic
                return text.strip()
            
        except Exception as e:
            logger.error(f"Error reading from database: {str(e)}", exc_info=True)
        
        # Fallback to file system
        try:
            filepath = os.path.join(Config.ARTICLES_DIR, 'docx', decoded_filename)
            normalized_path = os.path.normpath(filepath)
            
            if '..' in normalized_path:
                raise ValueError("Path traversal detected")
            
            if not os.path.exists(normalized_path):
                raise FileNotFoundError(f"Article file not found: {decoded_filename}")
            
            doc = Document(normalized_path)
            return "\n".join([para.text for para in doc.paragraphs])
            
        except Exception as e:
            logger.error(f"Error reading from file system: {str(e)}", exc_info=True)
            raise FileNotFoundError(f"Article file not found: {decoded_filename}")

    @staticmethod
    def read_markdown(filename):
        """
        Read content from a markdown file (database first, then file system fallback)
        Args:
            filename: Name of the markdown file (may be URL-encoded)
        Returns:
            Markdown content
        """
        # URL-decode the filename first
        decoded_filename = unquote(filename)
        
        try:
            # Try to read from database first
            db = get_db()
            cursor = db.cursor()
            cursor.execute("""
                SELECT markdown_content
                FROM articles 
                WHERE filename = ? AND is_active = 1 AND status = 'active'
            """, (decoded_filename,))
            article = cursor.fetchone()
            
            if article and (article[0] if not hasattr(article, 'keys') else article['markdown_content']):
                return article[0] if not hasattr(article, 'keys') else article['markdown_content']
            
        except Exception as e:
            logger.error(f"Error reading markdown from database: {str(e)}", exc_info=True)
        
        # Fallback to file system
        try:
            # Convert .docx extension to .md extension
            if decoded_filename.endswith('.docx'):
                markdown_filename = decoded_filename.replace('.docx', '.md')
            else:
                markdown_filename = decoded_filename
            
            filepath = os.path.join(Config.ARTICLES_DIR, 'markdown', markdown_filename)
            normalized_path = os.path.normpath(filepath)
            
            if '..' in normalized_path:
                raise ValueError("Path traversal detected")
            
            if not os.path.exists(normalized_path):
                raise FileNotFoundError(f"Markdown file not found: {markdown_filename}")
            
            with open(normalized_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"Error reading markdown from file system: {str(e)}", exc_info=True)
            raise FileNotFoundError(f"Markdown file not found: {decoded_filename}")
    
    @staticmethod
    def save_content(content):
        """
        Save generated content to a file
        Args:
            content: Content to save
        Returns:
            Filename of the saved content
        """
        filename = f"blog_{int(time.time())}.txt"
        path = os.path.join(Config.GENERATED_DIR, filename)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return filename

    @staticmethod
    def generate_formatted_docx(content, title="Legal Blog"):
        """Generate DOCX with exact formatting from markdown"""
        doc = Document()

        # Custom styles (can be modified)
        styles = {
            'h1': {'font_size': 16, 'bold': True, 'color': RGBColor(0, 32, 96)},
            'h2': {'font_size': 14, 'bold': True, 'color': RGBColor(0, 64, 128)},
            'h3': {'font_size': 12, 'bold': True, 'italic': True},
            'bold': {'bold': True},
            'normal': {'font_size': 11}
        }
        
        def apply_style(run, style):
            """Helper function to apply formatting"""
            run.font.size = Pt(style.get('font_size', 11))
            run.font.bold = style.get('bold', False)
            run.font.italic = style.get('italic', False)
            if 'color' in style:
                run.font.color.rgb = style['color']
        
        # Process markdown content line by line
        lines = content.split('\n')
        for line in lines:

            if line.replace('-', '').strip() == '' and len(line) >= 3:
                continue

            # Detect formatting
            if line.startswith('# '):  # H1
                para = doc.add_heading(level=1)
                run = para.add_run(line[2:].strip())
                apply_style(run, styles['h1'])
                
            elif line.startswith('## '):  # H2
                para = doc.add_heading(level=2)
                run = para.add_run(line[3:].strip())
                apply_style(run, styles['h2'])
                
            elif line.startswith('### '):  # H3
                para = doc.add_heading(level=3)
                run = para.add_run(line[4:].strip())
                apply_style(run, styles['h3'])
                
            elif '**' in line:  # Bold text
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        apply_style(run, styles['bold'])
                    else:
                        para.add_run(part)
            
            else:  # Normal paragraph
                para = doc.add_paragraph()
                run = para.add_run(line)
                apply_style(run, styles['normal'])
        # Collect all empty paragraphs
        empty_paragraphs = [p for p in doc.paragraphs if not p.text.strip()]

        # Remove each empty paragraph from the document
        for p in empty_paragraphs:
            p._element.getparent().remove(p._element)

        # Save to bytes buffer
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

azure_services = AzureServices()
image_generator = ImageGenerator()

@app.template_filter('markdown')
def markdown_filter(text):
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, 'html.parser')
    return str(soup)

def require_admin(f):
    """Decorator to require admin access"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Check if user is logged in
        if 'user' not in session:
            return redirect(url_for('login'))
        
        user_data = session['user']
        user_id = user_data.get('id')
        
        # Check admin status from session first
        if user_data.get('is_admin'):
            return f(*args, **kwargs)
        
        # Fallback: check database
        db = get_db()
        cursor = db.cursor()
        cursor.execute("SELECT is_admin FROM users WHERE id = ?", (user_id,))
        user = cursor.fetchone()
        
        if not user or not user['is_admin']:
            flash('Access denied. Admin privileges required.', 'error')
            return redirect(url_for('dashboard'))
        
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
def home():
    if not UserSession.get_current_user():
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/register', methods=['GET', 'POST'])
@limiter.limit("3 per minute")  # Limit registration attempts to 3 per minute per IP
def register():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        firm = request.form.get('firm', '').strip()
        location = request.form.get('location', '').strip()
        lawyer_name = request.form.get('lawyer_name', '').strip()
        state = request.form.get('state', '').strip()
        address = request.form.get('address', '').strip()
        planning_session = request.form.get('planning_session', '').strip()
        other_planning_session = request.form.get('other_planning_session', '').strip()
        discovery_call_link = request.form.get('discovery_call_link', '').strip()

        success, error = UserSession.register(email, password, firm, location, lawyer_name, state, 
                              address, planning_session, other_planning_session, discovery_call_link)
        if success:
            # Auto-login after registration
            UserSession.login(email, password)
            return redirect(url_for('dashboard'))
        
        return render_template('register.html', error=error or "Registration failed. Please try again.")
    
    return render_template('register.html')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    # Check if user is blocked
    if UserSession.is_user_blocked(user['id']):
        session.clear()
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Handle both form data and JSON requests
        if request.is_json:
            data = request.get_json()
            firm = data.get('firm', '')
            location = data.get('location', '')
            lawyer_name = data.get('lawyer_name', '')
            state = data.get('state', '')
            address = data.get('address', '')
            planning_session = data.get('planning_session', '')
            discovery_call_link = data.get('discovery_call_link', '')
            selected_tone = data.get('selected_tone', '')
            tone_description = data.get('tone_description', '')
            keywords = data.get('keywords', '')
        else:
            firm = request.form['firm']
            location = request.form['location']
            lawyer_name = request.form['lawyer_name']
            state = request.form['state']
            address = request.form.get('address', '')
            planning_session = request.form.get('planning_session', '')
            discovery_call_link = request.form.get('discovery_call_link', '')
            selected_tone = request.form.get('selected_tone', '')
            tone_description = request.form.get('tone_description', '')
            keywords = request.form.get('keywords', '')
        
        if UserSession.update_profile(user['username'], firm, location, lawyer_name, state, address, planning_session, "", discovery_call_link, selected_tone, tone_description, keywords):
            session['user']['firm'] = firm
            session['user']['location'] = location
            session['user']['lawyer_name'] = lawyer_name
            session['user']['state'] = state
            session['user']['address'] = address
            session['user']['planning_session'] = planning_session
            session['user']['discovery_call_link'] = discovery_call_link
            session['user']['selected_tone'] = selected_tone
            session['user']['tone_description'] = tone_description
            session['user']['keywords'] = keywords
            session.modified = True
            
            # Return JSON response for AJAX requests
            if request.is_json:
                return jsonify({
                    'success': True,
                    'firm': firm,
                    'location': location,
                    'lawyer_name': lawyer_name,
                    'state': state,
                    'address': address,
                    'planning_session': planning_session,
                    'discovery_call_link': discovery_call_link,
                    'selected_tone': selected_tone,
                    'tone_description': tone_description,
                    'keywords': keywords
                })
            
            return redirect(url_for('dashboard'))
        
        # Return JSON response for AJAX requests
        if request.is_json:
            return jsonify({'success': False, 'error': 'Update failed'}), 400
        
        return render_template('profile.html', error="Update failed", user=session['user'])
    
    return render_template('profile.html', user=session['user'])

@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("5 per minute")  # Limit login attempts to 5 per minute per IP
def login():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        
        # Basic validation
        if not email or not password:
            return render_template('login.html', error="Email and password are required")
        
        email_valid, email_error = validate_email(email)
        if not email_valid:
            return render_template('login.html', error=email_error)
        
        if len(password) > 128:
            return render_template('login.html', error="Invalid credentials")
        
        if UserSession.login(email, password):
            logger.info(f"User logged in successfully: {email}")
            return redirect(url_for('dashboard'))
        
        logger.warning(f"Failed login attempt for email: {email}")
        return render_template('login.html', error="Invalid credentials")
    return render_template('login.html')

@app.route('/forgot_password', methods=['GET', 'POST'])
@limiter.limit("3 per hour")  # Limit password reset requests to 3 per hour per IP
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        
        # Validate email
        if not email:
            return render_template('forgot_password.html', error="Email is required")
        
        email_valid, email_error = validate_email(email)
        if not email_valid:
            return render_template('forgot_password.html', error=email_error)
        
        # Check if user exists
        db = get_db()
        user = db.execute('SELECT * FROM users WHERE email = ?', (email,)).fetchone()
        
        if user:
            # Generate reset token
            token = secrets.token_urlsafe(32)
            expires = datetime.now() + timedelta(hours=24)
            
            # Store reset token in database (SQL Server compatible)
            # First, delete any existing tokens for this email
            db.execute('DELETE FROM password_resets WHERE email = ?', (email,))
            
            # Then insert the new token
            db.execute('''
                INSERT INTO password_resets (email, token, expires)
                VALUES (?, ?, ?)
            ''', (email, token, expires))
            db.commit()
            
            # In a real application, you would send an email here
            # For now, we'll just show a success message with the token
            reset_url = url_for('reset_password', token=token, _external=True)
            
            return render_template('forgot_password.html', 
                                 success=f"Password reset link sent! For demo purposes, here's the link: {reset_url}")
        else:
            return render_template('forgot_password.html', 
                                 error="If an account with that email exists, a reset link has been sent.")
    
    return render_template('forgot_password.html')

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
@limiter.limit("5 per hour")  # Limit password reset attempts to 5 per hour per IP
def reset_password(token):
    db = get_db()
    
    # Check if token is valid and not expired
    reset_record = db.execute('''
        SELECT * FROM password_resets 
        WHERE token = ? AND expires > ? AND used = 0
    ''', (token, datetime.now())).fetchone()
    
    if not reset_record:
        return render_template('reset_password.html', token=token,
                             error="Invalid or expired reset link. Please request a new one.")
    
    if request.method == 'POST':
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        if password != confirm_password:
            return render_template('reset_password.html', token=token,
                                 error="Passwords do not match.")
        
        # Validate password strength
        password_valid, password_error = validate_password(password)
        if not password_valid:
            return render_template('reset_password.html', token=token, error=password_error)
        
        # Hash password before storing
        hashed_password = generate_password_hash(password)
        db.execute('UPDATE users SET password = ? WHERE email = ?', 
                  (hashed_password, reset_record[1]))
        
        # Mark token as used
        db.execute('UPDATE password_resets SET used = 1 WHERE token = ?', (token,))
        db.commit()
        
        return render_template('reset_password.html', token=token,
                             success="Password reset successfully! You can now login with your new password.")
    
    return render_template('reset_password.html', token=token)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    # Check if user is blocked
    if UserSession.is_user_blocked(user['id']):
        session.clear()
        return redirect(url_for('login'))
    
    # Get articles and their metadata
    articles = FileManager.list_articles()
    metadata = FileManager.get_article_metadata()
    
    # Get unique series names
    series_list = set()
    for article in articles:
        meta = metadata.get(article, {})
        if 'series' in meta:
            series_list.add(meta['series'])
    series_list = sorted(series_list) if series_list else None

    # Combine standard tones with user's custom tones
    standard_tones = [
        ('Professional', 'Formal and business-like tone suitable for corporate audiences'),
        ('Friendly', 'Warm and approachable tone that builds rapport with readers'),
        ('Educational', 'Clear and informative tone designed to explain concepts')
    ]
    
    # Get user's custom tones from database
    custom_tones = UserSession.get_custom_tones(user['id'])
    all_tones = standard_tones + [(t['name'], t['description']) for t in custom_tones]
    
    # Convert to the format expected by the template
    tone_options = [t[0] for t in all_tones]
    tone_descriptions = {t[0]: t[1] for t in all_tones}
    
    return render_template('dashboard.html', 
                         user=user,
                         username=user['username'],
                         articles=articles,
                         metadata=metadata,
                         tone_options=tone_options,
                         tone_descriptions=tone_descriptions,
                         user_keywords=user.get('keywords', ''),
                         series_list=series_list)

@app.route('/add_tone', methods=['POST'])
@limiter.limit("10 per hour")  # Limit custom tone creation
def add_tone():
    user = UserSession.get_current_user()
    if not user:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    data = request.get_json() if request.is_json else request.form
    tone_name = data.get('tone_name', '').strip()
    tone_description = data.get('tone_description', '').strip()
    
    if not tone_name:
        return jsonify({'success': False, 'error': 'Tone name is required'}), 400
    
    if UserSession.add_custom_tone(user['id'], tone_name, tone_description):
        return jsonify({
            'success': True,
            'tone_name': tone_name,
            'tone_description': tone_description
        })
    
    return jsonify({'success': False, 'error': 'Tone with this name already exists'}), 400

@app.route('/submit_feedback', methods=['POST'])
@limiter.limit("5 per hour")  # Limit feedback submissions
def submit_feedback():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'User not authenticated'})
    
    try:
        message = request.form.get('message')
        contact_email = request.form.get('contact_email')
        
        # Validate required fields
        if not message:
            return jsonify({'success': False, 'message': 'Please provide a message'})
        
        # Set default values for removed fields
        feedback_type = 'general'
        priority = 'medium'
        subject = 'User Feedback'
        
        # Submit feedback using UserSession method
        if UserSession.submit_feedback(session['user']['id'], feedback_type, priority, subject, message, contact_email):
            return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
        else:
                    return jsonify({'success': False, 'message': 'An error occurred while submitting feedback'})
    
    except Exception as e:
        logger.error(f"Error submitting feedback: {str(e)}", exc_info=True)
        return jsonify({'success': False, 'message': 'An error occurred while submitting feedback'})



@app.route('/select/<article>', methods=['GET', 'POST'])
@limiter.limit("10 per hour")  # Limit content generation to 10 per hour per user/IP
async def select_article(article):
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    # Initialize variables that will be used in both GET and POST
    firm = ''
    location = ''
    
    if request.method == 'POST':
        # Use user's saved tone and keywords from their profile
        tone = user.get('selected_tone', 'Professional')
        tone_description = user.get('tone_description', '')
        
        # If no custom tone description is saved, use default descriptions
        if not tone_description:
            default_descriptions = {
                'Professional': 'Formal and business-like tone suitable for corporate audiences',
                'Friendly': 'Warm and approachable tone that builds rapport with readers',
                'Educational': 'Clear and informative tone designed to explain concepts'
            }
            tone_description = default_descriptions.get(tone, 'Formal and business-like tone suitable for corporate audiences')
            
        keywords = user.get('keywords', '')
        firm = user.get('firm', '')
        location = user.get('location', '')
        lawyer_name = user.get('lawyer_name', '')
        city = user.get('location', '')
        state = user.get('state', '')
        planning_session_name = user.get('planning_session', '15-minute discovery call')
        discovery_call_link = user.get('discovery_call_link', '')

        # Track activity start time
        start_time = time.time()
        
        # Call Azure Function for content generation
        function_url = f"{FUNCTION_APP_URL}/api/content_generator?code={FUNCTION_KEY}"
        
        # Log content generation request
        logger.debug(f"Content generation request - Article: {article}, Tone: {tone}, Firm: {firm}")
        logger.debug(f"Function URL: {function_url}")
        
        payload = {
            "original_text": FileManager.read_docx(article),
            "tone": tone,
            "tone_description": tone_description,
            "keywords": keywords,
            "firm_name": firm,
            "location": location,
            "lawyer_name": lawyer_name,
            "city": city,
            "state": state,
            "planning_session_name": planning_session_name,
            "discovery_call_link": discovery_call_link
        }
        
        logger.debug(f"Payload prepared with {len(payload)} items")
        
        if SIMULATE_OPENAI:
            logger.info("Using SIMULATE_OPENAI mode for content generation")
            await simulate_openai_call()
            blog_content = "Simulated blog content"
            
            # Log simulated activity
            UserActivityTracker.log_activity(
                user_id=user['id'],
                activity_type="content_generation",
                feature_name="AI Article Generation",
                api_endpoint="SIMULATED",
                request_payload_size=len(str(payload)),
                response_status=200,
                response_size=len(blog_content),
                processing_time_ms=int((time.time() - start_time) * 1000),
                success=True,
                additional_data=f"Article: {article}, Tone: {tone}, Keywords: {keywords}"
            )
        else:
            logger.info("Making HTTP request to Azure Function for content generation")
            try:
                import aiohttp
                async with aiohttp.ClientSession() as client_session:
                    logger.debug(f"Sending POST request to: {function_url}")
                    logger.debug(f"Request payload size: {len(str(payload))} characters")
                    
                    async with client_session.post(function_url, json=payload) as response:
                        logger.debug(f"Response status: {response.status}")
                        
                        response_text = await response.text()
                        logger.debug(f"Response text length: {len(response_text)}")
                        
                        if response.status != 200:
                            logger.error(f"Content generation failed with status {response.status}: {response_text[:500]}")
                            
                            # Log failed activity
                            UserActivityTracker.log_activity(
                                user_id=user['id'],
                                activity_type="content_generation",
                                feature_name="AI Article Generation",
                                api_endpoint=function_url,
                                request_payload_size=len(str(payload)),
                                response_status=response.status,
                                response_size=len(response_text),
                                processing_time_ms=int((time.time() - start_time) * 1000),
                                success=False,
                                error_message=response_text,
                                additional_data=f"Article: {article}, Tone: {tone}, Keywords: {keywords}"
                            )
                            
                            raise Exception(f"Function error: {response_text}")
                        
                        result = await response.json()
                        logger.debug(f"Parsed JSON result keys: {list(result.keys())}")
                        blog_content = result["content"]
                        logger.info(f"Content generation successful. Generated content length: {len(blog_content)}")
                        
                        # Log successful activity
                        UserActivityTracker.log_activity(
                            user_id=user['id'],
                            activity_type="content_generation",
                            feature_name="AI Article Generation",
                            api_endpoint=function_url,
                            request_payload_size=len(str(payload)),
                            response_status=response.status,
                            response_size=len(blog_content),
                            processing_time_ms=int((time.time() - start_time) * 1000),
                            success=True,
                            additional_data=f"Article: {article}, Tone: {tone}, Keywords: {keywords}"
                        )
                        
            except Exception as e:
                logger.error(f"Content generation exception: {str(e)}", exc_info=True)
                
                # Log exception activity
                UserActivityTracker.log_activity(
                    user_id=user['id'],
                    activity_type="content_generation",
                    feature_name="AI Article Generation",
                    api_endpoint=function_url,
                    request_payload_size=len(str(payload)),
                    response_status=0,
                    response_size=0,
                    processing_time_ms=int((time.time() - start_time) * 1000),
                    success=False,
                    error_message=str(e),
                    additional_data=f"Article: {article}, Tone: {tone}, Keywords: {keywords}"
                )
                
                raise
        
        # Save the generated content to a file
        filename = FileManager.save_content(blog_content)
        
        # Set up the session data for the review page (without image initially)
        session['current_post'] = {
            'original': article,
            'content': blog_content,
            'image': None,  # Image will be generated later when requested
            'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
            'tone': tone,
            'filename': filename
        }
        
        # Initialize chat history
        session['chat_history'] = [{
            'role': 'assistant',
            'content': blog_content,
            'content_is_blog': True,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        }]
        
        # Generate a unique session ID for the chat
        session['session_id'] = os.urandom(16).hex()
        
        return redirect(url_for('review'))
    
    tone_options = [
        'Professional',
        'Friendly',
        'Educational'
    ]
    
    tone_descriptions = {
        'Professional': 'Formal and business-like tone suitable for corporate audiences',
        'Friendly': 'Warm and approachable tone that builds rapport with readers',
        'Educational': 'Clear and informative tone designed to explain concepts'
    }
    
    return render_template('select.html',
                         article_name=article,
                         tone_options=tone_options,
                         tone_descriptions=tone_descriptions,
                         firm=firm,
                         location=location)

@app.route('/use_version', methods=['POST'])
def use_version():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    selected_content = request.form['content']
    
    session['current_post']['content'] = selected_content
    session.modified = True
    
    return redirect(url_for('finalize'))

@app.route('/finalize')
def finalize():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    post = session['current_post']
    filename = FileManager.save_content(post['content'])
    image_url = url_for('static', filename=f'generated/{post["image"]}') if post.get('image') else None
    
    return render_template('finalize.html', 
                         post=post,
                         filename=filename,
                         image_url=image_url)

@app.route('/review', methods=['GET', 'POST'])
@limiter.limit("30 per hour")  # Limit content editing to 30 per hour per user/IP
async def review():
    # Check if we have a filename parameter but no current_post in session
    filename = request.args.get('filename')
    if filename and 'current_post' not in session:
        # Try to load the content from the file
        try:
            # Use safe path validation
            filepath = get_safe_file_path(Config.GENERATED_DIR, filename)
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Set up the session data
            session['current_post'] = {
                'content': content,
                'filename': filename,
                'created': datetime.now().strftime("%Y-%m-%d %H:%M")
            }
            
            # Initialize chat history
            session['chat_history'] = [{
                'role': 'assistant',
                'content': content,
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            }]
            
            # Generate a unique session ID for the chat
            session['session_id'] = os.urandom(16).hex()
        except (ValueError, FileNotFoundError) as e:
            logger.error(f"Error loading file: {e}", exc_info=True)
            return redirect(url_for('dashboard'))
        except Exception as e:
            logger.error(f"Error loading file: {e}", exc_info=True)
            return redirect(url_for('dashboard'))
    
    # If we still don't have current_post in session, redirect to dashboard
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    post = session['current_post']
    
    if 'session_id' not in session:
        session['session_id'] = os.urandom(16).hex()
    
    if 'chat_history' not in session:
        session['chat_history'] = [{
            'role': 'assistant',
            'content': post['content'],
            'content_is_blog': True,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        }]
    
    if request.method == 'POST':
        if 'edit_message' in request.form:  # Chat-style editing
            user_message = request.form['edit_message']
            
            current_content = next(
                (msg['content'] for msg in reversed(session['chat_history']) 
                 if msg['content_is_blog']),
                post['content']
            )
            
            # Track activity start time
            start_time = time.time()
            
            # Call Azure Function for content editing
            function_url = f"{FUNCTION_APP_URL}/api/content_editor?code={FUNCTION_KEY}"
            payload = {
                "session_id": session['session_id'],
                "user_message": user_message,
                "current_content": current_content
            }
            
            logger.debug(f"Content Editor - Function URL: {function_url}")
            logger.debug(f"Content Editor - Payload keys: {list(payload.keys())}")
            logger.debug(f"Content Editor - User message: {user_message[:50]}...")
            logger.debug(f"Content Editor - Current content length: {len(current_content)}")
            
            if SIMULATE_OPENAI:
                logger.info("Content Editor - Using SIMULATE_OPENAI mode")
                await simulate_openai_call()
                edited_content = current_content  # Return current content for simulation
                
                # Log simulated activity
                UserActivityTracker.log_activity(
                    user_id=session['user']['id'],
                    activity_type="content_editing",
                    feature_name="AI Content Editing",
                    api_endpoint="SIMULATED",
                    request_payload_size=len(str(payload)),
                    response_status=200,
                    response_size=len(edited_content),
                    processing_time_ms=int((time.time() - start_time) * 1000),
                    success=True,
                    additional_data=f"User message: {user_message[:100]}..."
                )
            else:
                logger.info("Content Editor - Making HTTP request to Azure Function")
                try:
                    import aiohttp
                    async with aiohttp.ClientSession() as client_session:
                        logger.debug(f"Content Editor - Sending POST request to: {function_url}")
                        
                        async with client_session.post(function_url, json=payload) as response:
                            logger.debug(f"Content Editor - Response status: {response.status}")
                            
                            response_text = await response.text()
                            logger.debug(f"Content Editor - Response text length: {len(response_text)}")
                            
                            if response.status != 200:
                                logger.error(f"Content Editor - Error response (status {response.status}): {response_text[:500]}")
                                
                                # Log failed activity
                                UserActivityTracker.log_activity(
                                    user_id=session['user']['id'],
                                    activity_type="content_editing",
                                    feature_name="AI Content Editing",
                                    api_endpoint=function_url,
                                    request_payload_size=len(str(payload)),
                                    response_status=response.status,
                                    response_size=len(response_text),
                                    processing_time_ms=int((time.time() - start_time) * 1000),
                                    success=False,
                                    error_message=response_text,
                                    additional_data=f"User message: {user_message[:100]}..."
                                )
                                
                                raise Exception(f"Function error: {response_text}")
                            
                            result = await response.json()
                            logger.debug(f"Content Editor - Parsed JSON result keys: {list(result.keys())}")
                            edited_content = result["edited_content"]
                            logger.info(f"Content Editor - Edit successful. Edited content length: {len(edited_content)}")
                            
                            # Log successful activity
                            UserActivityTracker.log_activity(
                                user_id=session['user']['id'],
                                activity_type="content_editing",
                                feature_name="AI Content Editing",
                                api_endpoint=function_url,
                                request_payload_size=len(str(payload)),
                                response_status=response.status,
                                response_size=len(edited_content),
                                processing_time_ms=int((time.time() - start_time) * 1000),
                                success=True,
                                additional_data=f"User message: {user_message[:100]}..."
                            )
                            
                except Exception as e:
                    logger.error(f"Content Editor - Exception occurred: {str(e)}", exc_info=True)
                    
                    # Log exception activity
                    UserActivityTracker.log_activity(
                        user_id=session['user']['id'],
                        activity_type="content_editing",
                        feature_name="AI Content Editing",
                        api_endpoint=function_url,
                        request_payload_size=len(str(payload)),
                        response_status=0,
                        response_size=0,
                        processing_time_ms=int((time.time() - start_time) * 1000),
                        success=False,
                        error_message=str(e),
                        additional_data=f"User message: {user_message[:100]}..."
                    )
                    
                    raise
            
            session['chat_history'].append({
                'role': 'user',
                'content': user_message,
                'content_is_blog': False,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })
            session['chat_history'].append({
                'role': 'assistant',
                'content': edited_content,
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })
            
            post['content'] = edited_content
            session['current_post'] = post
            
        elif 'content' in request.form:  # Manual editing
            post['content'] = request.form['content']
            session['current_post'] = post
            session['chat_history'].append({
                'role': 'assistant',
                'content': post['content'],
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })

        session.modified = True
        return redirect(url_for('review'))
    
    # Save the current content to a file and get the filename
    if 'filename' not in post:
        filename = FileManager.save_content(post['content'])
        post['filename'] = filename
        session['current_post'] = post
    
    # Get source article content if available
    source_article_content = None
    if 'original' in post:
        try:
            markdown_content = FileManager.read_markdown(post['original'])
            # Convert markdown to HTML for proper display
            source_article_content = markdown.markdown(markdown_content)
        except Exception as e:
            logger.error(f"Error reading source article: {e}", exc_info=True)
            source_article_content = "Source article not available"
    
    image_url = url_for('static', filename=f'generated/{post["image"]}') if post.get('image') else None
    
    return render_template('review.html', 
                         post=post,
                         chat_history=session['chat_history'],
                         source_article_content=source_article_content,
                         image_url=image_url)

@app.route('/save_changes', methods=['POST'])
def save_changes():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    edited_content = request.form.get('content', '')
    
    session['current_post']['content'] = edited_content
    
    if 'chat_history' not in session:
        session['chat_history'] = []
    
    session['chat_history'].append({
        'role': 'system',
        'content': 'User saved manual changes',
        'content_is_blog': False,
        'timestamp': datetime.now().strftime("%H:%M:%S")
    })
    
    session.modified = True
    return redirect(url_for('finalize'))

@app.route('/download/<filename>')
def download(filename):
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    try:
        # For generated files, use simple path validation since they're internally generated
        filepath = os.path.join(Config.GENERATED_DIR, filename)
        normalized_path = os.path.normpath(filepath)
        
        # Simple path traversal check - look for .. in the path
        if '..' in normalized_path:
            logger.warning(f"Path traversal detected: {normalized_path}")
            return redirect(url_for('review'))
        
        # Check if file exists
        if not os.path.exists(normalized_path):
            logger.warning(f"File not found: {normalized_path}")
            return redirect(url_for('review'))
        
        with open(normalized_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Get title
        title = session['current_post'].get('original', 'Legal Blog').replace('.docx', '')
        
        # Generate formatted DOCX
        docx_file = FileManager.generate_formatted_docx(content, title)
        
        return send_file(
            docx_file,
            as_attachment=True,
            download_name=f"{title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Download error: {e}", exc_info=True)
        return redirect(url_for('review'))

@app.route('/generate_image')
@limiter.limit("20 per hour")  # Limit image generation to 20 per hour per user/IP
async def generate_image():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))

    # Track activity start time
    start_time = time.time()
    
    # Call Azure Function for image generation
    function_url = f"{FUNCTION_APP_URL}/api/image_generator?code={FUNCTION_KEY}"
    payload = {
        "text_prompt": session['current_post']['content']
    }
    
    logger.debug(f"Image Generator - Function URL: {function_url}")
    logger.debug(f"Image Generator - Payload keys: {list(payload.keys())}")
    logger.debug(f"Image Generator - Text prompt length: {len(payload['text_prompt'])}")
    
    if SIMULATE_OPENAI:
        logger.info("Image Generator - Using SIMULATE_OPENAI mode")
        await simulate_openai_call()
        session['current_post']['image'] = "dummy.png"
        session.modified = True
        
        # Log simulated activity
        UserActivityTracker.log_activity(
            user_id=session['user']['id'],
            activity_type="image_generation",
            feature_name="AI Image Generation",
            api_endpoint="SIMULATED",
            request_payload_size=len(str(payload)),
            response_status=200,
            response_size=0,
            processing_time_ms=int((time.time() - start_time) * 1000),
            success=True,
            additional_data="Generated dummy image"
        )
        
        return redirect(url_for('review'))
    
    logger.info("Image Generator - Making HTTP request to Azure Function")
    try:
        import aiohttp, base64, os
        async with aiohttp.ClientSession() as client_session:
            logger.debug(f"Image Generator - Sending POST request to: {function_url}")
            
            async with client_session.post(function_url, json=payload) as response:
                logger.debug(f"Image Generator - Response status: {response.status}")
                
                response_text = await response.text()
                logger.debug(f"Image Generator - Response text length: {len(response_text)}")
                
                if response.status != 200:
                    logger.error(f"Image Generator - Error response (status {response.status}): {response_text[:500]}")
                    
                    # Log failed activity
                    UserActivityTracker.log_activity(
                        user_id=session['user']['id'],
                        activity_type="image_generation",
                        feature_name="AI Image Generation",
                        api_endpoint=function_url,
                        request_payload_size=len(str(payload)),
                        response_status=response.status,
                        response_size=len(response_text),
                        processing_time_ms=int((time.time() - start_time) * 1000),
                        success=False,
                        error_message=response_text,
                        additional_data="Image generation failed"
                    )
                    
                    raise Exception(f"Function error: {response_text}")
                
                result = await response.json()
                logger.debug(f"Image Generator - Parsed JSON result keys: {list(result.keys())}")
                logger.info(f"Image Generator - Image generation successful. Filename: {result.get('image_filename', 'unknown')}")
                
                # Log successful activity
                UserActivityTracker.log_activity(
                    user_id=session['user']['id'],
                    activity_type="image_generation",
                    feature_name="AI Image Generation",
                    api_endpoint=function_url,
                    request_payload_size=len(str(payload)),
                    response_status=response.status,
                    response_size=len(response_text),
                    processing_time_ms=int((time.time() - start_time) * 1000),
                    success=True,
                    additional_data=f"Generated image: {result.get('image_filename', 'unknown')}"
                )
                
    except Exception as e:
        logger.error(f"Image Generator - Exception occurred: {str(e)}", exc_info=True)
        
        # Log exception activity
        UserActivityTracker.log_activity(
            user_id=session['user']['id'],
            activity_type="image_generation",
            feature_name="AI Image Generation",
            api_endpoint=function_url,
            request_payload_size=len(str(payload)),
            response_status=0,
            response_size=0,
            processing_time_ms=int((time.time() - start_time) * 1000),
            success=False,
            error_message=str(e),
            additional_data="Image generation exception"
        )
        
        raise
    
    image_filename = result["image_filename"]
    os.makedirs(os.path.join(app.static_folder, 'generated'), exist_ok=True)
    
    # Save the image data to file
    image_path = os.path.join(app.static_folder, 'generated', image_filename)
    with open(image_path, 'wb') as f:
        f.write(base64.b64decode(result["image_data"]))

    if image_filename:
        session['current_post']['image'] = image_filename
        session.modified = True
    
    return redirect(url_for('review'))
    
# Database cleanup is handled by close_db() function above

@app.route('/preview_article/<article>')
def preview_article(article):
    try:
        # URL-decode the article filename first
        decoded_article = unquote(article)
        
        # For article filenames, we need to be more permissive
        # Check if the file exists in the articles directory
        article_path = os.path.join(Config.ARTICLES_DIR, 'docx', decoded_article)
        normalized_article_path = os.path.normpath(article_path)
        
        # Simple path traversal check - look for .. in the path
        if '..' in normalized_article_path:
            return jsonify({'error': 'Invalid article path'}), 400
        
        # Try to read markdown file first
        markdown_filename = decoded_article.replace('.docx', '.md')
        markdown_path = os.path.join(Config.ARTICLES_DIR, 'markdown', markdown_filename)
        normalized_markdown_path = os.path.normpath(markdown_path)
        
        # Ensure markdown path is also safe
        if '..' in normalized_markdown_path:
            return jsonify({'error': 'Invalid markdown path'}), 400
        
        if os.path.exists(normalized_markdown_path):
            # Read the markdown content
            with open(normalized_markdown_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            # If markdown doesn't exist, read from docx
            if not os.path.exists(normalized_article_path):
                return jsonify({'error': 'Article not found'}), 404
                
            doc = Document(normalized_article_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            
        # Convert the content to HTML for preview
        html_content = markdown.markdown(content)
        return jsonify({'content': html_content})
    except (ValueError, FileNotFoundError) as e:
        logger.error(f"File access error in preview_article: {str(e)}", exc_info=True)
        return jsonify({'error': 'Article not found'}), 404
    except Exception as e:
        logger.error(f"Error in preview_article: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

def is_safe_filename(filename: str) -> bool:
    """
    Check if a filename is safe to use (prevents path traversal attacks).
    
    Args:
        filename: Filename to validate
        
    Returns:
        True if filename is safe, False otherwise
    """
    """Validate that filename is safe and doesn't contain path traversal characters"""
    if not filename:
        return False
    
    # Check for path traversal characters
    dangerous_chars = ['..', '/', '\\', ':', '*', '?', '"', '<', '>', '|']
    for char in dangerous_chars:
        if char in filename:
            return False
    
    # Check if filename is only alphanumeric, dots, hyphens, and underscores
    if not re.match(r'^[a-zA-Z0-9._-]+$', filename):
        return False
    
    return True

def get_safe_file_path(base_dir, filename):
    """Safely construct a file path within the base directory"""
    if not is_safe_filename(filename):
        raise ValueError("Invalid filename")
    
    # Normalize the path to prevent path traversal
    full_path = os.path.normpath(os.path.join(base_dir, filename))
    
    # Ensure the path is within the base directory
    if not full_path.startswith(os.path.abspath(base_dir)):
        raise ValueError("Path traversal detected")
    
    return full_path

@app.route('/refresh-session')
def refresh_session():
    """Refresh user session with latest data from database"""
    if 'user' not in session:
        return redirect(url_for('login'))
    
    user_id = session['user']['id']
    db = get_db()
    cursor = db.cursor()
    cursor.execute("SELECT * FROM users WHERE id = ?", (user_id,))
    user = cursor.fetchone()
    
    if user:
        # Get user's custom tones
        cursor.execute('SELECT name, description FROM tones WHERE user_id = ?', (user.id,))
        tones = cursor.fetchall()
        
        session['user'] = {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'firm': user.firm,
            'location': user.location,
            'lawyer_name': user.lawyer_name,
            'state': user.state,
            'address': user.address,
            'planning_session': user.planning_session,
            'other_planning_session': user.other_planning_session,
            'discovery_call_link': user.discovery_call_link,
            'is_admin': getattr(user, 'is_admin', False),
            'custom_tones': [{'name': tone.name, 'description': tone.description} for tone in tones]
        }
        session.modified = True
        flash('Session refreshed successfully!', 'success')
    
    return redirect(url_for('dashboard'))

@app.route('/admin/articles', methods=['GET', 'POST'])
@require_admin
def admin_articles():
    """Admin interface for managing articles"""
    db = get_db()
    cursor = db.cursor()
    
    if request.method == 'POST':
        # Handle file upload
        if 'article_file' not in request.files:
            flash('No file selected', 'error')
            return redirect(url_for('admin_articles'))
        
        file = request.files['article_file']
        title = request.form.get('title', '').strip()
        description = request.form.get('description', '').strip()
        
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(url_for('admin_articles'))
        
        if not title:
            flash('Title is required', 'error')
            return redirect(url_for('admin_articles'))
        
        if not file.filename.lower().endswith('.docx'):
            flash('Only DOCX files are allowed', 'error')
            return redirect(url_for('admin_articles'))
        
        try:
            # Read file content
            file_content = file.read()
            
            # Convert DOCX to Markdown using existing function
            from content.articles.docx_to_markdown import convert_docx_to_markdown
            import io
            
            # Create a temporary file-like object
            docx_io = io.BytesIO(file_content)
            docx_io.name = file.filename
            
            # Convert to markdown
            markdown_content = convert_docx_to_markdown(docx_io)
            
            # Store in database
            cursor.execute("""
                INSERT INTO articles (title, description, filename, markdown_content, docx_content, created_by)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (title, description, file.filename, markdown_content, file_content, session['user']['id']))
            
            db.commit()
            flash(f'Article "{title}" uploaded successfully!', 'success')
            
        except Exception as e:
            db.rollback()
            flash(f'Error uploading article: {str(e)}', 'error')
            logger.error(f"Upload error: {e}", exc_info=True)
    
    # Get all articles for display
    cursor.execute("""
        SELECT a.*, u.username as created_by_name
        FROM articles a
        LEFT JOIN users u ON a.created_by = u.id
        ORDER BY a.created_at DESC
    """)
    articles = cursor.fetchall()
    
    return render_template('admin/articles.html', articles=articles)

# Error handlers for standardized error handling
@app.errorhandler(404)
def not_found_error(error):
    logger.warning(f"404 error: {request.url}")
    return render_template('error.html', error_code=404, error_message="Page not found"), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500 error: {str(error)}", exc_info=True)
    return render_template('error.html', error_code=500, error_message="Internal server error"), 500

@app.errorhandler(403)
def forbidden_error(error):
    logger.warning(f"403 error: {request.url}")
    return render_template('error.html', error_code=403, error_message="Access forbidden"), 403

@app.errorhandler(400)
def bad_request_error(error):
    logger.warning(f"400 error: {request.url}")
    return render_template('error.html', error_code=400, error_message="Bad request"), 400

if __name__ == '__main__':
    # Debug mode should only be enabled in development
    # Set FLASK_DEBUG=true in .env for development, false/absent for production
    debug_mode = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
    app.run(host='0.0.0.0', port=8000, debug=debug_mode)